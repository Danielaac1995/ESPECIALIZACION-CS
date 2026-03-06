"""
Generador de FDE-088-Propuesta Daniel Aguirre.docx  (v2)
Sigue la estructura EXACTA del formulario FDE-088 Versión 06
Sectores: Ficha Técnica → Antecedentes → Problema → Objetivos
          → Metodología → Referencias (IEEE) → Cronograma → Compromiso
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT_DIR  = r"c:\Users\Usuario\Desktop\Especializacion CS\01_Materias\Seminario_1\TAREA1"
OUTPUT_FILE = os.path.join(OUTPUT_DIR, "FDE-088-Propuesta Daniel Aguirre.docx")

# ─────────────────────────────────────────────
# HELPERS
# ─────────────────────────────────────────────
def shd(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    s    = OxmlElement("w:shd")
    s.set(qn("w:val"),   "clear")
    s.set(qn("w:color"), "auto")
    s.set(qn("w:fill"),  hex_color)
    tcPr.append(s)

def page_header(doc: Document):
    """Encabezado institucional identico al PDF."""
    t = doc.add_table(rows=1, cols=3)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    c0, c1, c2 = t.cell(0, 0), t.cell(0, 1), t.cell(0, 2)
    shd(c0, "1F497D")
    c0.width = Cm(3)
    p0 = c0.paragraphs[0]; p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r0 = p0.add_run("ITM")
    r0.bold = True; r0.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF); r0.font.size = Pt(16)

    shd(c1, "FFFFFF")
    p1 = c1.paragraphs[0]; p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p1.add_run("PROPUESTA DE PROYECTO DE GRADO")
    r1.bold = True; r1.font.size = Pt(11); r1.font.name = "Arial"

    shd(c2, "D6E4F0")
    c2.width = Cm(4)
    for label, val in [("Código", "FDE 088"), ("Versión", "06"), ("Fecha", "24-02-2020")]:
        p2 = c2.add_paragraph() if c2.paragraphs[0].text else c2.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
        rb = p2.add_run(label + ": "); rb.bold = True; rb.font.size = Pt(8)
        rv = p2.add_run(val);          rv.font.size = Pt(8)

    doc.add_paragraph()

def section_title(doc: Document, num: str, title: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)
    text = f"{num}. {title.upper()}" if num else title.upper()
    r = p.add_run(text)
    r.bold = True; r.font.size = Pt(11); r.font.name = "Arial"
    return p

def body(doc: Document, text: str):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(4)
    for run in p.runs:
        run.font.size = Pt(11); run.font.name = "Times New Roman"
    return p

def bullet(doc: Document, text: str):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    r.font.size = Pt(11); r.font.name = "Times New Roman"
    p.paragraph_format.space_after = Pt(2)
    return p

def field_table(doc: Document, rows_data: list):
    t = doc.add_table(rows=len(rows_data), cols=2)
    t.style = "Table Grid"
    for i, (label, value) in enumerate(rows_data):
        lc, vc = t.cell(i, 0), t.cell(i, 1)
        shd(lc, "D6E4F0")
        lc.width = Cm(5.5)
        rl = lc.paragraphs[0].add_run(label)
        rl.bold = True; rl.font.size = Pt(10)
        vc.paragraphs[0].add_run(value).font.size = Pt(10)
    doc.add_paragraph()

# ─────────────────────────────────────────────
# DOCUMENTO
# ─────────────────────────────────────────────
doc = Document()

for s in doc.sections:
    s.top_margin    = Cm(2.5)
    s.bottom_margin = Cm(2.5)
    s.left_margin   = Cm(3)
    s.right_margin  = Cm(2.5)

doc.styles["Normal"].font.name = "Times New Roman"
doc.styles["Normal"].font.size = Pt(11)

# ══════════════════════════════════════════════
# PÁGINA 1 — FICHA TÉCNICA
# ══════════════════════════════════════════════
page_header(doc)
section_title(doc, "1", "FICHA TÉCNICA")

field_table(doc, [
    ("Modalidad:", "Trabajo de investigación aplicada"),
    ("Título:",
     "Fortalecimiento del modelo de ciberseguridad IaaS en nube híbrida para PYMES "
     "colombianas: integración de inteligencia de amenazas (CTI) y automatización "
     "de respuesta a incidentes (SOAR)"),
])

# Encabezados estudiante
t_eh = doc.add_table(rows=1, cols=3)
t_eh.style = "Table Grid"
for i, hdr in enumerate(["Nombre", "Cédula", "Correo electrónico"]):
    shd(t_eh.cell(0, i), "D6E4F0")
    r = t_eh.cell(0, i).paragraphs[0].add_run(hdr)
    r.bold = True; r.font.size = Pt(10)

# Valores estudiante
t_ev = doc.add_table(rows=1, cols=3)
t_ev.style = "Table Grid"
for i, val in enumerate(["Daniel Aguirre", "Por diligenciar", "daniel.aguirre@itm.edu.co"]):
    t_ev.cell(0, i).paragraphs[0].add_run(val).font.size = Pt(10)
doc.add_paragraph()

field_table(doc, [("Programa Académico:", "Especialización en Ciberseguridad — ITM")])

# Encabezados asesor
t_ah = doc.add_table(rows=1, cols=3)
t_ah.style = "Table Grid"
for i, hdr in enumerate(["Nombre", "Institución", "Correo electrónico"]):
    shd(t_ah.cell(0, i), "D6E4F0")
    r = t_ah.cell(0, i).paragraphs[0].add_run(hdr)
    r.bold = True; r.font.size = Pt(10)

# Valores asesor
t_av = doc.add_table(rows=1, cols=3)
t_av.style = "Table Grid"
for i, val in enumerate(["Por asignar", "ITM", "Por diligenciar"]):
    t_av.cell(0, i).paragraphs[0].add_run(val).font.size = Pt(10)
doc.add_paragraph()

pnota = doc.add_paragraph()
pnota.add_run(
    "Nota: La propuesta tiene una extensión máxima de 9 páginas (letra 11, interlineado sencillo): "
    "5 páginas ítems 2–6, más ficha técnica, cronograma y compromiso."
).font.size = Pt(9)
doc.add_page_break()

# ══════════════════════════════════════════════
# PÁGINA 2 — SECCIONES 2 A 6
# ══════════════════════════════════════════════
page_header(doc)

# ──────────────────────────────────────────────
# 2. ANTECEDENTES Y MARCO TEÓRICO
# ──────────────────────────────────────────────
section_title(doc, "2", "ANTECEDENTES Y MARCO TEÓRICO")

body(doc,
    "Las PYMES colombianas representan más del 99% del tejido empresarial del país y son el "
    "segmento que más ha acelerado la adopción de servicios en la nube, especialmente IaaS en "
    "entornos híbridos, impulsado por la reducción de costos y la tercerización de TI [1]. "
    "Sin embargo, esta migración amplía la superficie de ataque en organizaciones que carecen "
    "de equipos especializados en ciberseguridad."
)

body(doc,
    "Moncada García (2025) desarrolló un modelo de ciberseguridad para IaaS en nube híbrida "
    "para PYMES, basado en ISO/IEC 27001, NIST SP 800-53 y el marco CSA, validado en un "
    "caso de estudio real [2]. El modelo identifica controles para amenazas críticas (accesos "
    "no autorizados, DDoS, fuga de datos) y constituye el punto de partida de esta propuesta."
)

body(doc,
    "La literatura reciente demuestra que los modelos basados en controles estáticos son "
    "insuficientes frente a amenazas dinámicas. La integración de CTI (Cyber Threat Intelligence) "
    "y plataformas SOAR reduce el MTTD y MTTR entre 50–70% [3][4]. Herramientas open-source "
    "como OpenCTI, MISP, TheHive y Cortex hacen viable esta integración en PYMES sin costos "
    "adicionales significativos [5]."
)

# ──────────────────────────────────────────────
# 3. PLANTEAMIENTO DEL PROBLEMA
# ──────────────────────────────────────────────
section_title(doc, "3", "PLANTEAMIENTO DEL PROBLEMA")

body(doc,
    "El modelo de Moncada García (2025) presenta las siguientes brechas operativas que "
    "esta propuesta busca resolver:"
)

for b in [
    "Ausencia de CTI: controles estáticos sin inteligencia de amenazas en tiempo real.",
    "Sin SOAR: respuesta a incidentes manual, inviable para equipos de TI pequeños.",
    "Sin KPIs operativos: no define MTTD ni MTTR para medir desempeño en producción.",
    "Monitoreo continuo sin operacionalizar: no especifica cómo implementar SIEM/IDS de bajo costo.",
]:
    bullet(doc, b)

p_preg = doc.add_paragraph()
r_preg = p_preg.add_run(
    "Pregunta problema: ¿Cómo puede fortalecerse el modelo IaaS de ciberseguridad para PYMES "
    "colombianas integrando CTI y SOAR, de forma técnica y económicamente viable?"
)
r_preg.bold = True; r_preg.italic = True
r_preg.font.size = Pt(11); r_preg.font.name = "Times New Roman"
p_preg.paragraph_format.space_before = Pt(6); p_preg.paragraph_format.space_after = Pt(6)

# ──────────────────────────────────────────────
# 4. OBJETIVOS
# ──────────────────────────────────────────────
section_title(doc, "4", "OBJETIVOS")

p41 = doc.add_paragraph()
r41 = p41.add_run("4.1 General")
r41.bold = True; r41.font.size = Pt(11); r41.font.name = "Arial"

body(doc,
    "Fortalecer el modelo de ciberseguridad IaaS en nube híbrida para PYMES de Moncada García "
    "(2025), integrando CTI y SOAR para reducir el tiempo de detección y contención de amenazas "
    "en organizaciones con recursos tecnológicos limitados."
)

p42 = doc.add_paragraph()
r42 = p42.add_run("4.2 Específicos")
r42.bold = True; r42.font.size = Pt(11); r42.font.name = "Arial"

for e in [
    "Identificar las brechas operativas del modelo base respecto a detección en tiempo real y automatización.",
    "Evaluar herramientas open-source de CTI y SOAR aplicables a PYMES con restricciones presupuestarias.",
    "Diseñar la arquitectura de integración CTI-SOAR con playbooks automatizados por vector de ataque.",
    "Definir KPIs operativos (MTTD, MTTR, cobertura de controles) y un dashboard de monitoreo.",
    "Validar el modelo fortalecido en un caso de estudio real con una PYME colombiana en nube híbrida.",
]:
    bullet(doc, e)

# ──────────────────────────────────────────────
# 5. METODOLOGÍA PROPUESTA
# ──────────────────────────────────────────────
section_title(doc, "5", "METODOLOGÍA PROPUESTA")

body(doc,
    "Investigación aplicada con enfoque mixto. Las fases son: (1) Revisión sistemática de "
    "literatura; (2) Análisis de brechas del modelo base; (3) Diseño del modelo fortalecido "
    "según la opción seleccionada; (4) Implementación en laboratorio (AWS Free Tier / Azure "
    "DevTest); (5) Validación en PYME real; (6) Análisis comparativo MTTD/MTTR."
)

p_op = doc.add_paragraph()
p_op.add_run("Opciones de mejora al modelo base — Se proponen estas alternativas:").bold = True
p_op.runs[0].font.size = Pt(10)
doc.add_paragraph()

opciones = [
    ("A — CTI + SOAR Open-Source",
     "Integrar OpenCTI/MISP con TheHive/Cortex usando MITRE ATT&CK para mapear TTPs a los "
     "controles existentes. Herramientas gratuitas, viable sin inversión adicional."),
    ("B — Zero Trust + Microsegmentación",
     "Ampliar el modelo con verificación continua de identidades, microsegmentación de IaaS "
     "y acceso mínimo privilegiado. Refuerza arquitecturalmente los controles de acceso del modelo original."),
    ("C — SIEM + Detección de anomalías con ML",
     "Incorporar Wazuh o ELK Stack con módulos de ML para detectar comportamientos anómalos. "
     "Operacionaliza el monitoreo continuo que el modelo base no especifica."),
    ("D — Modelo de madurez progresivo (CMM-PYME)",
     "Diseñar 5 niveles de madurez que permitan avanzar incrementalmente en la implementación. "
     "Aporta un camino claro y medible ausente en el modelo base."),
    ("E — Combinación A + C ✔ (Recomendada)",
     "Integrar CTI/SOAR (Opción A) con SIEM de bajo costo (Opción C). Cubre las brechas más "
     "críticas manteniendo el costo dentro del alcance de PYMES colombianas."),
]

t_op = doc.add_table(rows=len(opciones) + 1, cols=2)
t_op.style = "Table Grid"
# Cabecera
shd(t_op.cell(0, 0), "1F497D"); shd(t_op.cell(0, 1), "1F497D")
for col, hdr in [(0, "Opción"), (1, "Descripción y valor aportado al modelo existente")]:
    rc = t_op.cell(0, col).paragraphs[0].add_run(hdr)
    rc.bold = True; rc.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF); rc.font.size = Pt(10)

# Filas
for i, (op, desc) in enumerate(opciones, 1):
    bg = "EBF3FB" if i % 2 == 0 else "FFFFFF"
    shd(t_op.cell(i, 0), bg); shd(t_op.cell(i, 1), bg)
    t_op.cell(i, 0).width = Cm(4.5)
    ro = t_op.cell(i, 0).paragraphs[0].add_run(op)
    ro.bold = True; ro.font.size = Pt(9.5)
    rd = t_op.cell(i, 1).paragraphs[0].add_run(desc)
    rd.font.size = Pt(9.5)

doc.add_paragraph()
doc.add_page_break()

# ══════════════════════════════════════════════
# PÁGINA 3 — REFERENCIAS + CRONOGRAMA
# ══════════════════════════════════════════════
page_header(doc)

# ──────────────────────────────────────────────
# 6. REFERENCIAS BIBLIOGRÁFICAS (IEEE)
# ──────────────────────────────────────────────
section_title(doc, "6", "REFERENCIAS BIBLIOGRÁFICAS")

refs = [
    "[1]  MinTIC, \"Índice de adopción de TIC y transformación digital en empresas colombianas,\" "
    "Ministerio de Tecnologías de la Información y las Comunicaciones, Bogotá, 2022. "
    "[En línea]. Disponible: https://www.mintic.gov.co",

    "[2]  A. F. Moncada García, \"Modelo de ciberseguridad aplicado a la infraestructura como "
    "servicio IaaS usada en la nube híbrida para pymes, con base en gestión de riesgos,\" "
    "Trabajo de grado — Maestría, Institución Universitaria ITM, 2025. "
    "Disponible: https://hdl.handle.net/20.500.12622/8029",

    "[3]  Gartner, \"Market Guide for Security Orchestration, Automation and Response Solutions,\" "
    "Gartner Research, 2023. [En línea]. Disponible: https://www.gartner.com",

    "[4]  D. Schlette, F. Böhm, M. Caselli y G. Pernul, \"Measuring and visualizing cyber threat "
    "intelligence quality,\" Int. J. Inf. Secur., vol. 20, pp. 21–38, 2021. "
    "DOI: 10.1007/s10207-020-00490-y",

    "[5]  W. Tounsi y H. Rais, \"A survey on technical threat intelligence in the age of "
    "sophisticated cyber attacks,\" Comput. Secur., vol. 72, pp. 212–233, 2018. "
    "DOI: 10.1016/j.cose.2017.09.001",

    "[6]  ISO/IEC, \"ISO/IEC 27001:2022 — Information security management systems,\" "
    "International Organization for Standardization, Ginebra, 2022.",

    "[7]  NIST, \"SP 800-53 Rev. 5: Security and Privacy Controls for Information Systems,\" "
    "2020. DOI: 10.6028/NIST.SP.800-53r5",

    "[8]  Cloud Security Alliance, \"Cloud Controls Matrix v4.0,\" CSA, 2023. "
    "Disponible: https://cloudsecurityalliance.org/research/cloud-controls-matrix",
]

for ref in refs:
    p = doc.add_paragraph(ref)
    p.paragraph_format.left_indent        = Cm(0.7)
    p.paragraph_format.first_line_indent  = Cm(-0.7)
    p.paragraph_format.space_after        = Pt(4)
    for run in p.runs:
        run.font.size = Pt(10); run.font.name = "Times New Roman"

doc.add_paragraph()

# ──────────────────────────────────────────────
# 7. CRONOGRAMA DE ACTIVIDADES
# ──────────────────────────────────────────────
section_title(doc, "7", "CRONOGRAMA DE ACTIVIDADES")

metas = [
    ("1", "Revisión sistemática de literatura"),
    ("2", "Análisis de brechas del modelo base"),
    ("3", "Diseño del modelo fortalecido (arquitectura CTI-SOAR)"),
    ("4", "Implementación en laboratorio (AWS / Azure DevTest)"),
    ("5", "Validación en caso de estudio real con PYME colombiana"),
    ("6", "Análisis de resultados y redacción del informe final"),
    ("7", "Entrega y sustentación"),
]

t_cr = doc.add_table(rows=len(metas) + 2, cols=8)
t_cr.style = "Table Grid"

# Fila 0: Mes (cabecera global)
for col in range(8):
    shd(t_cr.cell(0, col), "1F497D")
r_mes = t_cr.cell(0, 2).paragraphs[0].add_run("Mes")
r_mes.bold = True; r_mes.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF); r_mes.font.size = Pt(9)
r_m0 = t_cr.cell(0, 0).paragraphs[0].add_run("Metas")
r_m0.bold = True; r_m0.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF); r_m0.font.size = Pt(9)
r_d0 = t_cr.cell(0, 1).paragraphs[0].add_run("")
t_cr.cell(0, 1).width = Cm(7.5)

# Fila 1: numeración 1-6
for col in range(8):
    shd(t_cr.cell(1, col), "D6E4F0")
for col in range(2, 8):
    rn = t_cr.cell(1, col).paragraphs[0].add_run(str(col - 1))
    rn.bold = True; rn.font.size = Pt(9)
    t_cr.cell(1, col).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

# Filas de metas
for i, (num, desc) in enumerate(metas, 2):
    bg = "FFFFFF" if i % 2 == 0 else "F7F7F7"
    for col in range(8):
        shd(t_cr.cell(i, col), bg)
    tn = t_cr.cell(i, 0).paragraphs[0].add_run(num)
    tn.bold = True; tn.font.size = Pt(9)
    t_cr.cell(i, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    td = t_cr.cell(i, 1).paragraphs[0].add_run(desc)
    td.font.size = Pt(9)
    mes = i - 1   # meta 1→mes1, meta 2→mes2 ...
    if 2 <= mes + 1 <= 7:
        col_idx = mes + 1
        shd(t_cr.cell(i, col_idx), "BDD7EE")
        t_cr.cell(i, col_idx).paragraphs[0].add_run("●").font.size = Pt(9)
        t_cr.cell(i, col_idx).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
doc.add_page_break()

# ══════════════════════════════════════════════
# PÁGINAS 4-5 — COMPROMISO
# ══════════════════════════════════════════════
page_header(doc)
section_title(doc, "", "COMPROMISO PARA EL DESARROLLO DE TRABAJOS DE GRADO")

# Fecha
t_f = doc.add_table(rows=2, cols=4)
t_f.style = "Table Grid"
for col, lbl in enumerate(["Fecha", "DD", "MM", "AA"]):
    shd(t_f.cell(0, col), "D6E4F0")
    t_f.cell(0, col).paragraphs[0].add_run(lbl).bold = True
for col, val in enumerate(["", "01", "03", "2026"]):
    t_f.cell(1, col).paragraphs[0].add_run(val).font.size = Pt(10)
doc.add_paragraph()

field_table(doc, [
    ("Título del Trabajo de Grado:",
     "Fortalecimiento del modelo de ciberseguridad IaaS en nube híbrida para PYMES colombianas: "
     "integración de CTI y SOAR"),
    ("Modalidad:", "Trabajo de investigación aplicada"),
])

# Estudiante
body(doc, "Estudiante(s)")
t_est_c = doc.add_table(rows=3, cols=2)
t_est_c.style = "Table Grid"
for i, (lbl, val) in enumerate([
    ("Nombre completo:", "Daniel Aguirre"),
    ("Cédula de ciudadanía:", "Por diligenciar"),
    ("Programa:", "Especialización en Ciberseguridad"),
]):
    shd(t_est_c.cell(i, 0), "D6E4F0")
    t_est_c.cell(i, 0).paragraphs[0].add_run(lbl).bold = True
    t_est_c.cell(i, 1).paragraphs[0].add_run(val).font.size = Pt(10)
doc.add_paragraph()

# Asesor
body(doc, "Asesor(es)")
t_as_c = doc.add_table(rows=3, cols=2)
t_as_c.style = "Table Grid"
for i, (lbl, val) in enumerate([
    ("Nombre completo:", "Por asignar"),
    ("Cédula de ciudadanía:", "Por diligenciar"),
    ("Departamento:", "Facultad de Ingeniería — ITM"),
]):
    shd(t_as_c.cell(i, 0), "D6E4F0")
    t_as_c.cell(i, 0).paragraphs[0].add_run(lbl).bold = True
    t_as_c.cell(i, 1).paragraphs[0].add_run(val).font.size = Pt(10)
doc.add_paragraph()

# Compromisos asesor
body(doc, "El(Los) asesor(es) del trabajo de grado se compromete(n) con los siguientes deberes:")
for c in [
    "a) Orientar el desarrollo técnico y metodológico para cumplir los objetivos del trabajo.",
    "b) Hacer seguimiento del cronograma aprobado.",
    "c) Informar al Comité de Trabajos de Grado sobre cualquier anomalía en el desarrollo.",
    "d) Cumplir y velar por el respeto a las normas de propiedad intelectual y derechos de autor.",
]:
    p = doc.add_paragraph(c)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after  = Pt(2)
    for run in p.runs: run.font.size = Pt(10)

doc.add_paragraph()

# Compromisos estudiante
body(doc, "El(Los) estudiante(s) se compromete(n) con los siguientes deberes:")
for c in [
    "a) Manejar confidencialmente la información institucional y usarla exclusivamente para los fines del trabajo.",
    "b) Respetar la propiedad intelectual de terceros y evitar el plagio.",
    "c) Asumir la responsabilidad ante cualquier reclamación sobre derechos de autor.",
    "d) Abstenerse de cometer faltas disciplinarias como falsificación de firmas o suplantación.",
]:
    p = doc.add_paragraph(c)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after  = Pt(2)
    for run in p.runs: run.font.size = Pt(10)

doc.add_paragraph()
body(doc,
    "Los aspectos de propiedad intelectual se regirán por el Estatuto de Propiedad Intelectual "
    "del ITM, Acuerdo No. 34 de julio 23 de 2013 del Consejo Directivo."
)
body(doc, "En señal de aceptación se firma este documento.")
doc.add_paragraph()

# Firmas
t_firmas = doc.add_table(rows=2, cols=2)
t_firmas.style = "Table Grid"
shd(t_firmas.cell(0, 0), "D6E4F0"); shd(t_firmas.cell(0, 1), "D6E4F0")
t_firmas.cell(0, 0).paragraphs[0].add_run("FIRMA ESTUDIANTE(S)").bold = True
t_firmas.cell(0, 1).paragraphs[0].add_run("FIRMA ASESOR(ES)").bold = True
t_firmas.cell(1, 0).paragraphs[0].add_run(
    "\n\n_______________________________\nDaniel Aguirre"
).font.size = Pt(10)
t_firmas.cell(1, 1).paragraphs[0].add_run(
    "\n\n_______________________________\n(Por asignar)"
).font.size = Pt(10)

doc.add_paragraph()
doc.add_paragraph().add_run("FECHA ENTREGA: _______________").font.size = Pt(10)

# ══════════════════════════════════════════════
# PIE DE PÁGINA
# ══════════════════════════════════════════════
footer = doc.sections[0].footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
rf = footer.add_run(
    "FDE-088 v06  |  Especialización en Ciberseguridad — ITM  |  Daniel Aguirre  |  Marzo 2026"
)
rf.font.size = Pt(8); rf.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

# ══════════════════════════════════════════════
# GUARDAR
# ══════════════════════════════════════════════
doc.save(OUTPUT_FILE)
print(f"✅ Documento generado:\n   {OUTPUT_FILE}")
